from langchain_ollama import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
import os
import fitz 
import docx  
import pandas as pd 
import shutil
import spacy  
import time
from datetime import datetime
from config import Converted_TXT, Source_Documents , Vectors, embedding_model, Chunk_size

embeddings = OllamaEmbeddings(model=embedding_model)

def load_vectors():
    global vector_store
    vector_store = Chroma(
        collection_name="database",
        persist_directory=Vectors, 
        embedding_function=embeddings
    )
    return vector_store

vector_store = load_vectors()

def retrieve_database_from_vector(question):
    retriever = vector_store.as_retriever()
    docs = retriever.invoke(question)
    database_text = "\n".join([doc.page_content for doc in docs])
    return database_text

def read_pdf(file_path):
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text

def read_word(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def read_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_string() 

def convert_and_store(Source_Documents):
    for file_name in os.listdir(Source_Documents):
        file_path = os.path.join(Source_Documents, file_name)
        converted_path = os.path.join(Converted_TXT, f"{os.path.splitext(file_name)[0]}.txt")

        if os.path.exists(converted_path):
            continue

        if file_name.endswith(".pdf"):
            content = read_pdf(file_path)
        elif file_name.endswith(".docx"):
            content = read_word(file_path)
        elif file_name.endswith(".xlsx"):
            content = read_excel(file_path)
        elif file_name.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as txt_file:
                content = txt_file.read()
        else:
            continue

        with open(converted_path, "w", encoding="utf-8") as txt_file:
            txt_file.write(content)
    
    if os.path.exists(Source_Documents):
        for file_name in os.listdir(Source_Documents):
            file_path = os.path.join(Source_Documents, file_name)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.remove(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(f"Error deleting files: {file_path}: {e}")

#nlp = spacy.load("de_core_news_sm") # German Language Support
nlp = spacy.load("en_core_web_sm") # English Language Support

def split_text_into_chunks(text, max_tokens=Chunk_size):
    doc = nlp(text)
    sentences = [sentence.text.strip() for sentence in doc.sents if sentence.text.strip()]
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        combined_text = (current_chunk + " " + sentence).strip() if current_chunk else sentence
        token_count = len(combined_text.split())
        if token_count > max_tokens:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
        else:
            current_chunk = combined_text

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks

def txt_to_vector():
    converted_folder_path = os.path.join(os.path.expanduser("~"), Converted_TXT)

    documents = []

    for file_name in os.listdir(converted_folder_path):
        file_path = os.path.join(converted_folder_path, file_name)
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        text_chunks = split_text_into_chunks(text)
        for chunk in text_chunks:
            document = Document(
                page_content=chunk,
                metadata={"file_name": file_name}
            )
            documents.append(document)

    if documents:
        vector_store.add_documents(documents)

    if os.path.exists(converted_folder_path):
        try:
            for file_name in os.listdir(converted_folder_path):
                file_path = os.path.join(converted_folder_path, file_name)
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.remove(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
        except Exception as e:
            print(f"Error deleting files: {e}")
    print("Task completed successfully.")

last_run_hour = None

def sync_and_delete():
    global last_run_hour
    folder_path_local = os.path.join(os.path.expanduser("~"), Source_Documents)
    time_rn = datetime.now().hour
    if time_rn % 4 == 0 and time_rn != last_run_hour:
        convert_and_store(folder_path_local)
        txt_to_vector()
        load_vectors()
        print(f"Synchronization completed at {time_rn}:00.")
        last_run_hour = time_rn

